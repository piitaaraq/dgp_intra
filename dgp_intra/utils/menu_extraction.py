"""
Menu extraction utilities for dgp_intra.
Extracts menu data from Word documents for patient menus.
"""

from docx import Document
import re


def extract_patients_menu_from_docx(file_path):
    """
    Extract patient menu data from a Word document.
    
    Reads a standardized weekly menu document and extracts:
    - Week number
    - Lunch "Varmt" dishes for all 7 days
    - Dinner "Varmt" dishes for all 7 days
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        dict: Menu data structured as:
            {
                'week_number': int,
                'monday': {'lunch': str, 'dinner': str},
                'tuesday': {'lunch': str, 'dinner': str},
                ...
            }
            
    Raises:
        Exception: If document structure doesn't match expected format
    """
    doc = Document(file_path)
    
    if not doc.tables:
        raise ValueError("Ingen tabel fundet i dokumentet")
    
    table = doc.tables[0]
    
    # Extract week number from first cell (e.g., "Uge 50")
    week_text = table.rows[0].cells[0].text.strip()
    week_match = re.search(r'Uge (\d+)', week_text)
    week_number = int(week_match.group(1)) if week_match else None
    
    if not week_number:
        raise ValueError("Kunne ikke finde ugenummer i dokumentet")
    
    # Row indices for "Varmt" items
    LUNCH_VARMT_ROW = 2
    DINNER_VARMT_ROW = 5
    
    # Detect document format by checking number of columns
    num_columns = len(table.columns)
    
    if num_columns == 15:
        # Original kitchen format (merged cells)
        day_columns = {
            'monday': 2,
            'tuesday': 4,
            'wednesday': 6,
            'thursday': 8,
            'friday': 10,
            'saturday': 12,
            'sunday': 13,
        }
    elif num_columns == 8:
        # Generated format (clean 8 columns)
        day_columns = {
            'monday': 1,
            'tuesday': 2,
            'wednesday': 3,
            'thursday': 4,
            'friday': 5,
            'saturday': 6,
            'sunday': 7,
        }
    else:
        raise ValueError(f"Uventet tabel format: {num_columns} kolonner. Forventet 8 eller 15.")
    
    menu_data = {'week_number': week_number}
    
    # Extract data for each day
    for day_name, col_idx in day_columns.items():
        try:
            lunch_text = table.rows[LUNCH_VARMT_ROW].cells[col_idx].text.strip()
            dinner_text = table.rows[DINNER_VARMT_ROW].cells[col_idx].text.strip()
        except IndexError:
            raise ValueError(f"Kunne ikke læse data for {day_name}")
        
        # Clean up the text
        # Remove "Varmt:" label if present
        lunch_text = lunch_text.replace('Varmt:', '').strip()
        dinner_text = dinner_text.replace('Varmt:', '').strip()
        
        # Remove "Smørrebrød" if it appears at the end (part of "Koldt" section)
        dinner_text = re.sub(r'\n\s*Smørrebrød\s*$', '', dinner_text).strip()
        
        menu_data[day_name] = {
            'lunch': lunch_text if lunch_text else None,
            'dinner': dinner_text if dinner_text else None,
        }
    
    return menu_data