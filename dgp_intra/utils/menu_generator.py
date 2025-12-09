"""
Generate Word document for patient menu in the kitchen's standard format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def rgb_to_hex(rgb):
    """Convert RGB tuple to hex string."""
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])


def lighten_color(hex_color, factor):
    """
    Lighten a hex color by a factor (0.0 to 1.0).
    factor=0.5 means 50% lighter (towards white)
    """
    rgb = hex_to_rgb(hex_color)
    # Move each component towards 255 (white)
    new_rgb = tuple(int(c + (255 - c) * factor) for c in rgb)
    return rgb_to_hex(new_rgb)


def generate_color_scheme(base_color):
    """
    Generate a color scheme from a base color.
    Returns dict with header, section, and alternating row colors.
    """
    return {
        'header': base_color,           # Base color (darkest)
        'section': lighten_color(base_color, 0.4),  # 40% lighter
        'alt_row': lighten_color(base_color, 0.7),  # 70% lighter
    }


def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_vertical_alignment(cell, align="center"):
    """
    Set vertical alignment for cell.
    align can be: "top", "center", "bottom"
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def generate_patient_menu_docx(menu_data, output_path, base_color='4472C4'):
    """
    Generate a Word document for patient menu.
    
    Args:
        menu_data: Dictionary with structure:
            {
                'week_number': int,
                'monday': {'lunch': str, 'dinner': str},
                'tuesday': {...},
                ...
            }
        output_path: Path where to save the .docx file
        base_color: Hex color string (without #) for the header. Default is blue '4472C4'
    """
    doc = Document()
    
    # Set landscape orientation and narrow margins
    sections = doc.sections
    for section in sections:
        section.orientation = 1  # 1 = landscape, 0 = portrait
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"Det Grønlandske Patienthjem\nMenuplan Uge {menu_data['week_number']}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Calibri'
    
    doc.add_paragraph()  # Spacing
    
    # Create table: 7 rows (header + Frokost section + Aften section) x 8 columns (label + 7 days)
    table = doc.add_table(rows=7, cols=8)
    table.style = 'Table Grid'
    table.autofit = False  # Disable autofit
    table.allow_autofit = False
    
    # Generate color scheme from base color
    colors = generate_color_scheme(base_color)
    HEADER_COLOR = colors['header']      # Darkest - for header row
    SECTION_COLOR = colors['section']    # Medium - for section headers
    ALT_ROW_LIGHT = colors['alt_row']    # Lightest - for alternating rows
    
    # Days of the week
    days = ['Mandag', 'Tirsdag', 'Onsdag', 'Torsdag', 'Fredag', 'Lørdag', 'Søndag']
    day_keys = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']
    
    # Set row heights to use ~80% of page height
    # Row 0: Header (days) - 0.5"
    # Rows 1, 4: Section headers (Frokost/Aften) - 0.35"
    # Rows 2, 3, 5, 6: Content (Varmt/Koldt) - 1.0"
    table.rows[0].height = Inches(0.5)   # Days header
    table.rows[1].height = Inches(0.35)  # Frokost header
    table.rows[2].height = Inches(1.0)   # Lunch Varmt
    table.rows[3].height = Inches(1.0)   # Lunch Koldt
    table.rows[4].height = Inches(0.35)  # Aften header
    table.rows[5].height = Inches(1.0)   # Dinner Varmt
    table.rows[6].height = Inches(1.0)   # Dinner Koldt
    
    # Row 0: Days header (TALLER with BIGGER FONT and darker color)
    table.rows[0].cells[0].text = f"Uge {menu_data['week_number']}"
    for i, day in enumerate(days, start=1):
        cell = table.rows[0].cells[i]
        cell.text = day
        # Bold, centered, and LARGER font with white text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(13)  # Increased from 12
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        set_cell_background(cell, HEADER_COLOR)
        set_cell_vertical_alignment(cell, "center")
    
    # First cell styling
    cell = table.rows[0].cells[0]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(13)  # Increased from 12
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    set_cell_background(cell, HEADER_COLOR)
    set_cell_vertical_alignment(cell, "center")
    
    # Row 1: "Frokost" header (section header color)
    table.rows[1].cells[0].text = "Frokost"
    cell = table.rows[1].cells[0]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Calibri'
    set_cell_background(cell, SECTION_COLOR)
    set_cell_vertical_alignment(cell, "center")
    # Apply to all cells in row
    for i in range(1, 8):
        set_cell_background(table.rows[1].cells[i], SECTION_COLOR)
        set_cell_vertical_alignment(table.rows[1].cells[i], "center")
    
    # Row 2: Lunch - Varmt (white background)
    table.rows[2].cells[0].text = "Varmt:"
    set_cell_vertical_alignment(table.rows[2].cells[0], "center")
    for i, day_key in enumerate(day_keys, start=1):
        cell = table.rows[2].cells[i]
        lunch_text = menu_data.get(day_key, {}).get('lunch', '')
        cell.text = lunch_text or ''
        set_cell_vertical_alignment(cell, "center")
        # Centered, smaller font, sans-serif
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)  # Increased from 9
                run.font.name = 'Calibri'
    
    # Row 3: Lunch - Koldt (alternating light color)
    table.rows[3].cells[0].text = "Koldt:"
    set_cell_background(table.rows[3].cells[0], ALT_ROW_LIGHT)
    set_cell_vertical_alignment(table.rows[3].cells[0], "center")
    for i in range(1, 8):
        cell = table.rows[3].cells[i]
        cell.text = "Smørrebrød\n\nSalat\n\nFrisk frugt"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)  # Increased from 9
                run.font.name = 'Calibri'
        set_cell_background(cell, ALT_ROW_LIGHT)
        set_cell_vertical_alignment(cell, "center")
    
    # Row 4: "Aften" header (section header color)
    table.rows[4].cells[0].text = "Aften"
    cell = table.rows[4].cells[0]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Calibri'
    set_cell_background(cell, SECTION_COLOR)
    set_cell_vertical_alignment(cell, "center")
    for i in range(1, 8):
        set_cell_background(table.rows[4].cells[i], SECTION_COLOR)
        set_cell_vertical_alignment(table.rows[4].cells[i], "center")
    
    # Row 5: Dinner - Varmt (white background)
    table.rows[5].cells[0].text = "Varmt:"
    set_cell_vertical_alignment(table.rows[5].cells[0], "center")
    for i, day_key in enumerate(day_keys, start=1):
        cell = table.rows[5].cells[i]
        dinner_text = menu_data.get(day_key, {}).get('dinner', '')
        cell.text = dinner_text or ''
        set_cell_vertical_alignment(cell, "center")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)  # Increased from 9
                run.font.name = 'Calibri'
    
    # Row 6: Dinner - Koldt (alternating light color)
    table.rows[6].cells[0].text = "Koldt:"
    set_cell_background(table.rows[6].cells[0], ALT_ROW_LIGHT)
    set_cell_vertical_alignment(table.rows[6].cells[0], "center")
    for i in range(1, 8):
        cell = table.rows[6].cells[i]
        cell.text = "Salat\n\nFrisk frugt"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)  # Increased from 9
                run.font.name = 'Calibri'
        set_cell_background(cell, ALT_ROW_LIGHT)
        set_cell_vertical_alignment(cell, "center")
    
    # Set EXACT column widths to fit within margins (10 inches available)
    table.columns[0].width = Inches(0.9)   # Label: 0.9"
    for i in range(1, 8):
        table.columns[i].width = Inches(1.3)  # Days: 7 × 1.3" = 9.1"
    # Total: 0.9 + 9.1 = 10.0 inches (exactly fits the page)
    
    # Save document
    doc.save(output_path)


def generate_from_patients_menu_model(patients_menu, output_path, base_color='4472C4'):
    """
    Generate document from PatientsMenu database model.
    
    Args:
        patients_menu: PatientsMenu model instance
        output_path: Path where to save the .docx file
        base_color: Hex color string (without #) for the header. Default is blue '4472C4'
    """
    # Extract week number from week string (e.g., "2024-W50" -> 50)
    week_number = int(patients_menu.week.split('-W')[1])
    
    menu_data = {
        'week_number': week_number,
        'monday': {
            'lunch': patients_menu.monday,
            'dinner': patients_menu.monday_dinner
        },
        'tuesday': {
            'lunch': patients_menu.tuesday,
            'dinner': patients_menu.tuesday_dinner
        },
        'wednesday': {
            'lunch': patients_menu.wednesday,
            'dinner': patients_menu.wednesday_dinner
        },
        'thursday': {
            'lunch': patients_menu.thursday,
            'dinner': patients_menu.thursday_dinner
        },
        'friday': {
            'lunch': patients_menu.friday,
            'dinner': patients_menu.friday_dinner
        },
        'saturday': {
            'lunch': patients_menu.saturday,
            'dinner': patients_menu.saturday_dinner
        },
        'sunday': {
            'lunch': patients_menu.sunday,
            'dinner': patients_menu.sunday_dinner
        }
    }
    
    generate_patient_menu_docx(menu_data, output_path, base_color)
